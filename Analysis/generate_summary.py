#!/usr/bin/env python3
"""
=============================================================================
  Generate Summary Document for Co-Worker
  "Does It Pay to Play Nice?"
  Applied Sports Research — League of Legends

  Produces a Word document (Summary_for_Coworker.docx) containing:
    - Research questions & hypotheses
    - Data collection methodology
    - Variable operationalization (incl. early-game controls)
    - Descriptive statistics
    - Model results (baseline + extended + robustness)
    - Embedded figures
    - Reverse causality discussion
    - Notes on remaining work

  Usage:
      cd Analysis/
      python generate_summary.py
=============================================================================
"""

import numpy as np
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

BASE    = Path(__file__).parent
FIG_DIR = BASE / "figures"
TBL_DIR = BASE / "tables"
OUT_DOC = BASE.parent / "Summary_for_Coworker.docx"
DATA_CSV = BASE.parent / "Data Collection" / "match_dataset.csv"

doc = Document()

# ─── Styles ──────────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
font = style.font
font.name = "Helvetica"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_figure(path, caption, width=Inches(5.8)):
    """Insert a figure with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Normal"]
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
    return cap


def add_excel_table(xlsx_path, caption, max_rows=30):
    """Read an Excel file and insert it as a Word table."""
    df = pd.read_excel(xlsx_path, index_col=0)
    df = df.head(max_rows)
    for col in df.columns:
        if df[col].dtype == float:
            df[col] = df[col].round(4)

    table = doc.add_table(rows=1 + len(df), cols=1 + len(df.columns))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    for j, col in enumerate(df.columns):
        hdr[j + 1].text = str(col)

    for i, (idx, row) in enumerate(df.iterrows()):
        cells = table.rows[i + 1].cells
        cells[0].text = str(idx)
        for j, val in enumerate(row):
            cells[j + 1].text = str(val) if pd.notna(val) else ""

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    return table


# =============================================================================
#  TITLE PAGE
# =============================================================================

doc.add_paragraph("")
doc.add_paragraph("")
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Does It Pay to Play Nice?")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "The Effect of Toxic Behavior and Leadership-Style Communication\n"
    "on Win Rates in League of Legends")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph("")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    "Working Summary Document\n"
    "Applied Sports Research — University of Zurich\n"
    "Gian Senpinar | February 2026\n\n"
    "Status: Data collection + analysis complete. "
    "Ready for full paper drafting.")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# =============================================================================
#  TABLE OF CONTENTS (manual)
# =============================================================================

doc.add_heading("Contents", level=1)
toc_items = [
    "1. Research Questions & Hypotheses",
    "2. Data Collection Methodology",
    "3. Variable Operationalization",
    "4. Addressing Reverse Causality (Philippe Feedback)",
    "5. Descriptive Statistics",
    "6. Model Results",
    "   6.1 Baseline Model (Original Controls)",
    "   6.2 Extended Model (+ Early-Game Controls + Skill Proxy)",
    "   6.3 Model Comparison",
    "   6.4 Classification Performance",
    "7. Robustness Checks",
    "   7.1 Per-Region Sub-Regressions",
    "   7.2 Close Games Only",
    "8. Visualizations",
    "9. Remaining Work for Full Paper",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =============================================================================
#  1. RESEARCH QUESTIONS & HYPOTHESES
# =============================================================================

doc.add_heading("1. Research Questions & Hypotheses", level=1)

doc.add_paragraph(
    "This study investigates the relationship between in-game behavioral "
    "patterns and competitive outcomes in League of Legends, a team-based "
    "multiplayer online game. Using data from the Riot Games Match-V5 API, "
    "we examine whether toxic behavior and leadership-style communication "
    "are associated with win probability, controlling for team skill and "
    "early-game state.")

doc.add_heading("Research Questions", level=2)

rqs = [
    ("RQ1 (Toxicity)", "Is higher team-level toxic behavior associated with "
     "lower win probability, after controlling for team performance and "
     "early-game state?"),
    ("RQ2 (Leadership)", "Is leadership-style communication (structured ping "
     "usage, objective coordination) independently associated with higher "
     "win probability?"),
    ("RQ3 (Moderation)", "Does the leadership effect depend on the level of "
     "toxicity within the team (interaction effect)?"),
]
for label, text in rqs:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)

doc.add_heading("Hypotheses", level=2)

hypotheses = [
    ("H1", "Teams exhibiting higher toxicity scores will have lower win "
     "probability (negative association), even after controlling for "
     "early-game state and team skill."),
    ("H2", "Teams with higher leadership communication scores will have "
     "higher win probability (positive association), independent of team "
     "skill and game state."),
    ("H3", "The positive effect of leadership is attenuated (moderated) "
     "at high toxicity levels, suggesting that toxic environments undermine "
     "coordination benefits."),
]
for label, text in hypotheses:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)

doc.add_paragraph(
    "Note: Following feedback from the supervisor (Philippe), results are "
    "framed as associative rather than strictly causal. Early-game controls "
    "are included to strengthen the interpretation, but the cross-sectional "
    "observational design cannot fully rule out reverse causality.")

# =============================================================================
#  2. DATA COLLECTION
# =============================================================================

doc.add_page_break()
doc.add_heading("2. Data Collection Methodology", level=1)

doc.add_paragraph(
    "All data was collected from the Riot Games Match-V5 API between "
    "March 2025 and February 2026. The collection pipeline is fully "
    "automated and resumable.")

doc.add_heading("Sampling Design", level=2)

specs = [
    ("Regions", "EUW, NA, KR, VN (4 major server regions)"),
    ("Rank Tiers", "Gold, Platinum, Emerald (mid-to-high ranked play)"),
    ("Queue Type", "Ranked Solo/Duo (Queue ID 420)"),
    ("Time Window", "March 2025 -- February 2026 (12 months)"),
    ("Stratification", "Quarterly: equal quota per quarter to ensure "
     "temporal balance"),
    ("Validation Filters", "Game duration >= 15 min, no remakes, no early "
     "disconnects (<5 min played)"),
    ("Target", "~2,500 valid matches per region"),
]
for label, val in specs:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(val)

doc.add_heading("Final Sample", level=2)

df = pd.read_csv(DATA_CSV, keep_default_na=False)
n_matches = len(df) // 2
n_rows = len(df)

sample_table = doc.add_table(rows=6, cols=4)
sample_table.style = "Light Grid Accent 1"
headers = ["Region", "Matches", "Observations (2 per match)", "Date Range"]
for j, h in enumerate(headers):
    sample_table.rows[0].cells[j].text = h

regions_data = []
for region in ["EUW", "KR", "NA", "VN"]:
    rdf = df[df["region"] == region]
    regions_data.append((
        region, len(rdf) // 2, len(rdf),
        f"{rdf['game_date'].min()} to {rdf['game_date'].max()}"))

for i, (reg, m, obs, dr) in enumerate(regions_data):
    row = sample_table.rows[i + 1]
    row.cells[0].text = reg
    row.cells[1].text = str(m)
    row.cells[2].text = str(obs)
    row.cells[3].text = dr

total_row = sample_table.rows[5]
total_row.cells[0].text = "Total"
total_row.cells[1].text = str(n_matches)
total_row.cells[2].text = str(n_rows)
total_row.cells[3].text = ""
for cell in total_row.cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

doc.add_paragraph(
    f"\nThe final dataset contains {n_matches:,} unique matches "
    f"({n_rows:,} team-level observations) across {len(df.columns)} "
    f"variables.")

# =============================================================================
#  3. VARIABLE OPERATIONALIZATION
# =============================================================================

doc.add_page_break()
doc.add_heading("3. Variable Operationalization", level=1)

doc.add_heading("Dependent Variable", level=2)
doc.add_paragraph(
    "Win (binary): 1 if the team won the match, 0 otherwise. "
    "One row per team per match (2 rows per match in the full dataset). "
    "For regression analysis, we use the blue-side perspective (team_id = 100), "
    "yielding one observation per match.")

doc.add_heading("IV1: Toxicity Score", level=2)
doc.add_paragraph(
    "A standardized composite of four behavioral indicators of toxic play. "
    "Each component is z-scored across the full sample, then summed:")

tox_components = [
    ("max_feeding_index", "Highest Deaths/(Kills+Assists+1) on the team. "
     "Captures intentional feeding or extreme underperformance."),
    ("early_surrender", "Binary: 1 if the losing team surrendered. "
     "Proxy for team tilt / giving up."),
    ("vision_neglect_score", "Proportion of the team with < 0.3 vision/min. "
     "Proxy for disengagement / griefing."),
]
for var, desc in tox_components:
    p = doc.add_paragraph()
    run = p.add_run(f"{var}: ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(desc)

doc.add_heading("IV2: Leadership Score", level=2)
doc.add_paragraph(
    "A standardized composite of four leadership/coordination indicators. "
    "Each component is z-scored, then summed:")

lead_components = [
    ("coord_ping_ratio", "Top-pinging player's coordinating pings "
     "(onMyWay, command, assistMe, push) as a share of their total pings. "
     "Shot-caller proxy."),
    ("objectives_taken", "Total dragons + barons + rift heralds taken. "
     "Objective coordination."),
    ("vision_leadership", "Binary: 1 if both Support (>= 1.0 vis/min) and "
     "Jungle (>= 0.6 vis/min) meet vision thresholds."),
]
for var, desc in lead_components:
    p = doc.add_paragraph()
    run = p.add_run(f"{var}: ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(desc)

doc.add_heading("Interaction Term", level=2)
doc.add_paragraph(
    "toxicity_x_leadership = toxicity_score * leadership_score. "
    "Tests whether leadership effects are moderated by toxicity level.")

doc.add_heading("Control Variables (Original)", level=2)
controls_desc = [
    ("team_avg_kda", "Team-average (Kills + Assists) / max(Deaths, 1)"),
    ("team_avg_cs_min", "Team-average CS per minute"),
    ("team_gold", "Total gold earned by team"),
    ("team_damage", "Total damage dealt to champions"),
    ("team_vision_score", "Sum of vision scores across 5 players"),
    ("first_blood", "Binary: 1 if team got first blood"),
    ("game_duration_min", "Game length in minutes"),
    ("region dummies", "EUW (reference), KR, NA, VN"),
]
for var, desc in controls_desc:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{var}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Early-Game Controls (New, addressing reverse causality)", level=2)
doc.add_paragraph(
    "These variables capture the game state before toxicity and "
    "coordination behaviors fully manifest, providing pre-treatment "
    "controls:")

early_controls = [
    ("team_early_cs_10", "Team sum of lane + jungle CS at 10 minutes. "
     "Captures early laning competence independent of outcome."),
    ("team_early_gold_adv", "Team average early laning phase gold/XP "
     "advantage. Controls for which team was ahead early."),
    ("team_early_takedowns", "Team sum of takedowns in the first minutes. "
     "Captures early aggression / skirmish success."),
]
for var, desc in early_controls:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{var}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Skill Proxy (New)", level=2)
doc.add_paragraph(
    "avg_skill_index: Team average of (soloKills + skillshotsHit + "
    "skillshotsDodged) per player. This mechanical skill proxy controls "
    "for the confound that better players both coordinate more and win "
    "more, as highlighted in Philippe's feedback.")

# =============================================================================
#  4. ADDRESSING REVERSE CAUSALITY
# =============================================================================

doc.add_page_break()
doc.add_heading("4. Addressing Reverse Causality (Supervisor Feedback)", level=1)

doc.add_paragraph(
    "Philippe's central feedback identified endogeneity as the key "
    "methodological challenge:")

doc.add_heading("The Problem", level=2)

problems = [
    ("H1 (Toxicity)", "Toxicity may be a consequence of losing, not a "
     "cause. If teams that fall behind become more toxic, we would observe "
     "a negative toxicity-win correlation even without a causal effect."),
    ("H2 (Coordination)", "Teams that are winning can more easily secure "
     "objectives, place vision, and use strategic pings. The observed "
     "positive coordination-win correlation may reflect game advantage "
     "rather than independent leadership effects."),
    ("Player Skill", "Better players both coordinate more AND win more. "
     "Without controlling for skill, the leadership effect may be "
     "confounded."),
]
for label, text in problems:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)

doc.add_heading("Our Mitigation Strategy", level=2)

strategies = [
    ("1. Early-game controls", "By including team_early_cs_10, "
     "team_early_gold_adv, and team_early_takedowns, we control for the "
     "game state before mid-to-late-game behaviors (toxicity, coordination) "
     "fully manifest. If the IV effects persist after conditioning on "
     "early-game state, the association is less likely to be purely driven "
     "by reverse causality."),
    ("2. Skill proxy", "avg_skill_index captures mechanical skill "
     "(solo kills, skillshot accuracy) independent of team coordination. "
     "Including it separates individual ability from team-level "
     "communication patterns."),
    ("3. Close-games robustness", "Restricting the sample to games with "
     "< 5,000 gold difference removes blowouts where game state dominates. "
     "If leadership effects persist in close games, the association is "
     "more plausible as an independent factor."),
    ("4. Associative framing", "Results are presented as associations, "
     "not causal effects. The discussion should explicitly acknowledge "
     "remaining endogeneity concerns."),
]
for label, text in strategies:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)

# =============================================================================
#  5. DESCRIPTIVE STATISTICS
# =============================================================================

doc.add_page_break()
doc.add_heading("5. Descriptive Statistics", level=1)

desc_path = TBL_DIR / "B1_descriptive_stats.xlsx"
if desc_path.exists():
    add_excel_table(desc_path,
                    "Table 1: Descriptive Statistics (all variables)")

doc.add_paragraph("")

if (FIG_DIR / "B2_score_distributions.png").exists():
    add_figure(FIG_DIR / "B2_score_distributions.png",
               "Figure 1: Toxicity and Leadership Score Distributions by Region")

doc.add_paragraph("")

if (FIG_DIR / "B2_component_distributions.png").exists():
    add_figure(FIG_DIR / "B2_component_distributions.png",
               "Figure 2: IV Component Distributions")

doc.add_paragraph("")

if (FIG_DIR / "B3_outlier_boxplots.png").exists():
    add_figure(FIG_DIR / "B3_outlier_boxplots.png",
               "Figure 3: Outlier Box Plots by Region")

doc.add_paragraph("")

if (FIG_DIR / "B4_correlation_heatmap.png").exists():
    add_figure(FIG_DIR / "B4_correlation_heatmap.png",
               "Figure 4: Correlation Matrix (IVs, Controls incl. Early-Game, DV)")

doc.add_heading("Multicollinearity (VIF)", level=2)

vif_path = TBL_DIR / "B5_vif.xlsx"
if vif_path.exists():
    add_excel_table(vif_path, "Table 2: Variance Inflation Factors")

doc.add_paragraph(
    "Note: team_gold and game_duration_min show elevated VIFs (>10), "
    "which is expected since longer games produce more gold. The key IVs "
    "(toxicity_score, leadership_score) have VIFs below 2.0, indicating "
    "no problematic multicollinearity for the variables of interest. "
    "The early-game controls all have VIFs below 2.1.")

# =============================================================================
#  6. MODEL RESULTS
# =============================================================================

doc.add_page_break()
doc.add_heading("6. Model Results", level=1)

doc.add_heading("6.1 Baseline Model (Original Controls)", level=2)

doc.add_paragraph(
    "The baseline model includes the three IV composite scores, the "
    "original 7 control variables, and region dummies. This corresponds "
    "to the original analysis before Philippe's feedback.")

base_path = TBL_DIR / "C2_logistic_baseline.xlsx"
if base_path.exists():
    add_excel_table(base_path, "Table 3: Baseline Logistic Regression Results")

doc.add_heading("6.2 Extended Model (+ Early-Game Controls + Skill Proxy)",
                level=2)

doc.add_paragraph(
    "The extended model adds the four new variables addressing reverse "
    "causality: team_early_cs_10, team_early_gold_adv, "
    "team_early_takedowns, and avg_skill_index.")

ext_path = TBL_DIR / "C2_logistic_extended.xlsx"
if ext_path.exists():
    add_excel_table(ext_path, "Table 4: Extended Logistic Regression Results")

doc.add_heading("6.3 Model Comparison", level=2)

comp_path = TBL_DIR / "C2_model_comparison.xlsx"
if comp_path.exists():
    add_excel_table(comp_path,
                    "Table 5: Model Comparison (Key IVs: Baseline vs Extended)")

doc.add_paragraph("")

# Read actual comparison data for interpretation
if comp_path.exists():
    comp_df = pd.read_excel(comp_path, index_col=0)
    doc.add_heading("Key Findings from Model Comparison", level=3)

    findings = []

    if "toxicity_score" in comp_df.index:
        base_p = comp_df.loc["toxicity_score", "Baseline_p"]
        ext_p = comp_df.loc["toxicity_score", "Extended_p"]
        base_or = comp_df.loc["toxicity_score", "Baseline_OR"]
        ext_or = comp_df.loc["toxicity_score", "Extended_OR"]
        findings.append(
            f"Toxicity Score: OR = {base_or:.4f} (baseline, p={base_p:.4f}) "
            f"vs. OR = {ext_or:.4f} (extended, p={ext_p:.4f}). "
            f"The effect remains non-significant in both models, suggesting "
            f"that the null finding for toxicity is robust to the inclusion "
            f"of early-game controls.")

    if "leadership_score" in comp_df.index:
        base_p = comp_df.loc["leadership_score", "Baseline_p"]
        ext_p = comp_df.loc["leadership_score", "Extended_p"]
        base_or = comp_df.loc["leadership_score", "Baseline_OR"]
        ext_or = comp_df.loc["leadership_score", "Extended_OR"]
        findings.append(
            f"Leadership Score: OR = {base_or:.4f} (baseline, p={base_p:.4f}) "
            f"vs. OR = {ext_or:.4f} (extended, p={ext_p:.4f}). "
            f"The leadership effect remains highly significant and slightly "
            f"increases in magnitude with early-game controls, strengthening "
            f"the interpretation that the association is not merely driven "
            f"by game state.")

    if "toxicity_x_leadership" in comp_df.index:
        base_p = comp_df.loc["toxicity_x_leadership", "Baseline_p"]
        ext_p = comp_df.loc["toxicity_x_leadership", "Extended_p"]
        findings.append(
            f"Interaction (Toxicity x Leadership): The interaction term is "
            f"not significant in either model (baseline p={base_p:.4f}, "
            f"extended p={ext_p:.4f}).")

    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

doc.add_heading("6.4 Classification Performance", level=2)

doc.add_paragraph(
    "The extended model achieves strong predictive performance, though "
    "this is expected given that end-of-game statistics (gold, KDA, damage) "
    "are highly correlated with winning. The key interest is not prediction "
    "but the direction and significance of the IV coefficients after "
    "controlling for confounds.")

# =============================================================================
#  7. ROBUSTNESS CHECKS
# =============================================================================

doc.add_page_break()
doc.add_heading("7. Robustness Checks", level=1)

doc.add_heading("7.1 Per-Region Sub-Regressions (Extended Model)", level=2)

doc.add_paragraph(
    "The extended model was estimated separately for each region to test "
    "whether effects are consistent across cultural and competitive "
    "contexts:")

rob_path = TBL_DIR / "C4_robustness_by_region.xlsx"
if rob_path.exists():
    add_excel_table(rob_path,
                    "Table 6: Per-Region Robustness (Extended Model)")

doc.add_heading("7.2 Close Games Only (Gold Diff < 5,000)", level=2)

doc.add_paragraph(
    "To further address reverse causality, we restrict the sample to "
    "close games where the final gold difference between teams was less "
    "than 5,000. This removes blowout games where the losing team's "
    "behavior is likely a consequence of the game state. If the leadership "
    "effect persists in close games, it is more plausible as an independent "
    "factor.")

close_path = TBL_DIR / "C5_close_games_robustness.xlsx"
if close_path.exists():
    add_excel_table(close_path,
                    "Table 7: Close Games Robustness (gold diff < 5,000)")

# =============================================================================
#  8. VISUALIZATIONS
# =============================================================================

doc.add_page_break()
doc.add_heading("8. Visualizations", level=1)

figures = [
    ("D1_winrate_by_quartile.png",
     "Figure 5: Win Rate by Toxicity and Leadership Quartiles"),
    ("D2_region_comparison.png",
     "Figure 6: Regional Comparison of Toxicity, Leadership, and Pings"),
    ("D3_quarterly_trends.png",
     "Figure 7: Quarterly Trends in Toxicity and Leadership Scores"),
    ("D4_odds_ratio_forest.png",
     "Figure 8: Odds Ratio Forest Plot (Extended Model)"),
    ("D5_interaction_heatmap.png",
     "Figure 9: Toxicity x Leadership Interaction Heatmap"),
    ("D6_model_comparison.png",
     "Figure 10: Model Comparison (Baseline vs Extended vs Close Games)"),
]
for fname, caption in figures:
    path = FIG_DIR / fname
    if path.exists():
        add_figure(path, caption)
        doc.add_paragraph("")

# =============================================================================
#  9. REMAINING WORK
# =============================================================================

doc.add_page_break()
doc.add_heading("9. Remaining Work for Full Paper", level=1)

doc.add_paragraph(
    "The data collection, variable construction, and statistical analysis "
    "are complete. The following sections need to be written for the full "
    "seminar paper:")

remaining = [
    ("Theoretical Framework (Lazear)",
     "Philippe's feedback specifically asks for a clear connection to "
     "Lazear's theoretical framework. The paper should explicitly state "
     "which mechanism from Lazear is being tested and how the variables "
     "operationalize it. Current theories in the README (Social Identity "
     "Theory, Transformational Leadership, Team Conflict Theory) can be "
     "kept as supplementary, but Lazear should be the primary anchor. "
     "Key question: Which specific Lazear mechanism (tournament theory, "
     "personnel economics, incentive structures) maps to our toxicity/"
     "leadership constructs?"),
    ("Literature Review",
     "Expand the current references into a full literature review covering: "
     "(a) Lazear and personnel economics in competitive settings, "
     "(b) Toxicity in online games (Kokkinakis et al. 2020, Kwak et al. "
     "2015, Achterbosch et al. 2021), "
     "(c) Leadership and coordination in virtual teams, "
     "(d) Esports as a research context."),
    ("Methods Section",
     "Formalize the methodology: sampling design, variable construction "
     "(use the codebook), model specification with formula, and "
     "justification for logistic regression with interaction terms. "
     "Explicitly discuss the early-game controls as a strategy for "
     "addressing endogeneity."),
    ("Results Section",
     "Present the results in sequence: descriptive statistics, baseline "
     "model, extended model, comparison, robustness checks. All tables "
     "and figures are ready to embed."),
    ("Discussion",
     "Interpret findings in light of Lazear's framework. Key points: "
     "(1) Leadership shows a robust positive association with winning "
     "even after early-game controls and in close games; "
     "(2) Toxicity shows no significant independent association; "
     "(3) The interaction is not significant. "
     "Discuss reverse causality limitations explicitly. "
     "Discuss practical implications for game design, team management, "
     "and esports coaching."),
    ("Limitations",
     "Key limitations to discuss: "
     "(1) No chat data available (behavioral proxies only); "
     "(2) Cross-sectional design cannot establish causality; "
     "(3) No player-level rank/MMR data for skill control (proxy used); "
     "(4) Endogeneity mitigated but not eliminated; "
     "(5) Sample restricted to Gold-Emerald (generalizability); "
     "(6) team_gold and game_duration have high VIFs."),
]
for label, text in remaining:
    doc.add_heading(label, level=2)
    doc.add_paragraph(text)

# =============================================================================
#  APPENDIX: File Inventory
# =============================================================================

doc.add_page_break()
doc.add_heading("Appendix: File Inventory", level=1)

doc.add_heading("Scripts", level=2)
scripts = [
    ("Data Collection/collect_matches.py",
     "Riot API data collection (4 regions, stratified by quarter)"),
    ("Data Collection/build_variables.py",
     "Variable construction from raw JSONs, incl. early-game controls"),
    ("Analysis/analysis.py",
     "Full analysis pipeline (diagnostics + baseline/extended models + robustness)"),
    ("Analysis/generate_codebook.py",
     "Variable codebook generator"),
    ("Analysis/generate_summary.py",
     "This summary document generator"),
]
for path, desc in scripts:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{path}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Data Files", level=2)
data_files = [
    ("Data Collection/match_dataset.csv",
     f"{n_rows:,} rows x {len(df.columns)} columns (semicolon-delimited)"),
    ("Data Collection/raw_matches/",
     f"{n_matches:,} JSON files across 4 region subdirectories"),
    ("Multi-Region Dataset/variable_codebook.xlsx",
     "Complete variable documentation"),
]
for path, desc in data_files:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{path}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Tables (Analysis/tables/)", level=2)
for f in sorted(TBL_DIR.glob("*.xlsx")):
    doc.add_paragraph(f.name, style="List Bullet")

doc.add_heading("Figures (Analysis/figures/)", level=2)
for f in sorted(FIG_DIR.glob("*.png")):
    doc.add_paragraph(f.name, style="List Bullet")

# =============================================================================
#  SAVE
# =============================================================================

doc.save(str(OUT_DOC))
print(f"Summary document saved: {OUT_DOC}")
print(f"  Pages: ~{len(doc.paragraphs) // 30} (estimate)")
